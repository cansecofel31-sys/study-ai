from flask import Flask, render_template, request, redirect, session
from pypdf import PdfReader
import os

import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from flask import send_file
from docx import Document

texto_global = ""
app = Flask(__name__)
app.secret_key = "study_ai_admin_seguro"

def crear_base_datos():
    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            correo TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            plan TEXT DEFAULT 'gratis',
            fecha_registro TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS pagos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER,
            nombre TEXT,
            correo TEXT,
            plan TEXT,
            monto TEXT,
            comprobante TEXT,
            estado TEXT DEFAULT 'pendiente',
            fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS configuracion_pago (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            banco TEXT,
            titular TEXT,
            clabe TEXT,
            whatsapp TEXT,
            correo TEXT
        )
    """)

    cursor.execute("SELECT COUNT(*) FROM configuracion_pago")
    existe = cursor.fetchone()[0]

    if existe == 0:
        cursor.execute("""
            INSERT INTO configuracion_pago 
            (banco, titular, clabe, whatsapp, correo)
            VALUES (?, ?, ?, ?, ?)
        """, ("TU BANCO", "TU NOMBRE", "XXXX XXXX XXXX XXXX XX", "TU WHATSAPP", "contacto@studyai.com"))

    conexion.commit()
    conexion.close()

app.secret_key = "study_ai_clave_secreta"

def crear_base_datos():
    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            correo TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            plan TEXT DEFAULT 'gratis',
            fecha_registro TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS documentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER,
            nombre_archivo TEXT,
            resumen TEXT,
            total_palabras INTEGER,
            fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    conexion.commit()
    conexion.close()

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def limite_plan(plan):
    if plan == "gratis":
        return 3
    elif plan == "premium":
        return 50
    elif plan == "pro":
        return 999999
    return 3


def documentos_hoy(usuario_id):
    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        SELECT COUNT(*) FROM documentos
        WHERE usuario_id = ?
        AND DATE(fecha) = DATE('now')
    """, (usuario_id,))

    total = cursor.fetchone()[0]
    conexion.close()
    return total


def extraer_texto_pdf(ruta_pdf):
    texto = ""

    reader = PdfReader(ruta_pdf)

    for pagina in reader.pages:
        contenido = pagina.extract_text()
        if contenido:
            texto += contenido + "\n"

    return texto


def generar_resumen(texto):
    palabras = texto.split()

    if len(palabras) <= 600:
        return texto

    inicio = palabras[:200]
    medio = palabras[len(palabras)//2:len(palabras)//2 + 200]
    final = palabras[-200:]

    return " ".join(inicio + medio + final)


def generar_preguntas(texto):
    palabras = texto.split()
    tema = "el documento"

    if len(palabras) > 10:
        tema = " ".join(palabras[:5])

    return [
        f"¿Cuál es la idea principal de {tema}?",
        f"¿Qué conceptos importantes aparecen en {tema}?",
        f"¿Cuál es el objetivo principal de {tema}?",
        f"¿Qué información clave se debe recordar de {tema}?",
        f"¿Qué conclusión puedes obtener de {tema}?",
        "¿Qué conceptos aparecen con mayor frecuencia?",
        "¿Qué tema podría venir en un examen?",
        "¿Qué partes del documento son más importantes?",
        "¿Qué datos relevantes se presentan?",
        "¿Cómo se relaciona este tema con otros conocimientos?"
    ]

import random

def generar_flashcards(texto):

    preguntas = [

        "¿Cuál es la idea principal de esta sección del documento?",

        "¿Qué temas se desarrollan en esta parte del documento?",

        "¿Cuáles son las etapas del desarrollo del sistema?",

        "¿Qué conceptos importantes se explican aquí?",

        "¿Qué beneficios se mencionan en esta sección?",

        "¿Qué resultados se obtuvieron?",

        "¿Qué problema intenta resolver esta parte del documento?",

        "¿Qué conclusión puede obtenerse de esta sección?"

    ]

    palabras = texto.split()

    flashcards = []

    for i in range(0, min(len(palabras), 160), 20):

        bloque = palabras[i:i+20]

        flashcards.append({

            "pregunta": preguntas[len(flashcards) % len(preguntas)],

            "respuesta": " ".join(bloque)

        })

    return flashcards

@app.route("/", methods=["GET", "POST"])
def index():

    if "usuario_id" not in session:
        return redirect("/login")

    resumen = ""
    preguntas = []
    flashcards = []
    total_palabras = 0
    nombre_archivo = ""
    mensaje = ""

    if request.method == "POST":

        usuario_id = session["usuario_id"]
        plan = session.get("plan", "gratis")

        usados = documentos_hoy(usuario_id)
        limite = limite_plan(plan)

        if usados >= limite:
            mensaje = f"Has alcanzado tu límite diario del plan {plan}."
            return render_template(
                "index.html",
                resumen=resumen,
                preguntas=preguntas,
                flashcards=flashcards,
                total_palabras=total_palabras,
                nombre_archivo=nombre_archivo,
                mensaje=mensaje
            )

        archivo = request.files["pdf"]

        if archivo.filename != "":
            nombre_archivo = archivo.filename

            ruta = os.path.join(
                app.config["UPLOAD_FOLDER"],
                archivo.filename
            )

            archivo.save(ruta)

            texto = extraer_texto_pdf(ruta)

            total_palabras = len(texto.split())
            resumen = generar_resumen(texto)
            preguntas = generar_preguntas(texto)
            flashcards = generar_flashcards(texto)

    return render_template(
        "index.html",
        resumen=resumen,
        preguntas=preguntas,
        flashcards=flashcards,
        total_palabras=total_palabras,
        nombre_archivo=nombre_archivo,
        mensaje=mensaje
    )
@app.route("/descargar_pdf")
def descargar_pdf():
    archivo_pdf = "resumen_studyia.pdf"

    doc = SimpleDocTemplate(archivo_pdf)
    estilos = getSampleStyleSheet()

    contenido = [
        Paragraph("Resumen generado por Study IA", estilos["Title"]),
        Paragraph("Este es el PDF de prueba de Study IA.", estilos["BodyText"])
    ]

    doc.build(contenido)

    return send_file(archivo_pdf, as_attachment=True)

@app.route("/descargar_word")
def descargar_word():
    archivo_word = "cuestionario_studyia.docx"

    doc = Document()
    doc.add_heading("Cuestionario generado por Study IA", 0)

    preguntas = [
        "¿Cuál es la idea principal del documento?",
        "¿Qué conceptos importantes aparecen en el documento?",
        "¿Cuál es el objetivo principal del documento?",
        "¿Qué información clave se debe recordar?",
        "¿Qué conclusión puedes obtener?",
        "¿Qué conceptos aparecen con mayor frecuencia?",
        "¿Qué tema podría venir en un examen?",
        "¿Qué partes del documento son más importantes?",
        "¿Qué datos relevantes se presentan?",
        "¿Cómo se relaciona este tema con otros conocimientos?"
    ]

    for i, pregunta in enumerate(preguntas, start=1):
        doc.add_paragraph(f"{i}. {pregunta}")

    doc.save(archivo_word)

    return send_file(archivo_word, as_attachment=True)

@app.route("/preguntar", methods=["POST"])
def preguntar():
    pregunta = request.form.get("pregunta", "").lower()

    if not pregunta:
        return {"respuesta": "Escribe una pregunta primero."}

    uploads = os.listdir(app.config["UPLOAD_FOLDER"])

    if not uploads:
        return {"respuesta": "Primero carga un PDF."}

    ultimo_pdf = os.path.join(app.config["UPLOAD_FOLDER"], uploads[-1])
    texto = extraer_texto_pdf(ultimo_pdf)

    palabras = pregunta.split()

    for linea in texto.split("\n"):
        for palabra in palabras:
            if len(palabra) > 4 and palabra in linea.lower():
                return {"respuesta": linea[:800]}

    return {"respuesta": "No encontré información relacionada en el PDF."}

@app.route("/mapa_conceptual")
def mapa_conceptual():
    from PIL import Image, ImageDraw, ImageFont

    archivo = "mapa_conceptual_studyia.png"

    img = Image.new("RGB", (1200, 800), "#0f172a")
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype("arial.ttf", 40)
        font_box = ImageFont.truetype("arial.ttf", 28)
    except:
        font_title = ImageFont.load_default()
        font_box = ImageFont.load_default()

    draw.text((350, 40), "Mapa Conceptual - Study IA", fill="white", font=font_title)

    cajas = {
        "Documento PDF": (450, 140, 750, 210),
        "Ideas principales": (100, 330, 400, 400),
        "Conceptos clave": (800, 330, 1100, 400),
        "Resumen": (100, 560, 400, 630),
        "Preguntas": (450, 560, 750, 630),
        "Flashcards": (800, 560, 1100, 630),
    }

    for texto, box in cajas.items():
        draw.rounded_rectangle(box, radius=20, fill="#1e293b", outline="#38bdf8", width=3)
        draw.text((box[0] + 35, box[1] + 20), texto, fill="white", font=font_box)

    lineas = [
        ((600, 210), (250, 330)),
        ((600, 210), (950, 330)),
        ((250, 400), (250, 560)),
        ((950, 400), (600, 560)),
        ((950, 400), (950, 560)),
    ]

    for inicio, fin in lineas:
        draw.line([inicio, fin], fill="#38bdf8", width=4)

    img.save(archivo)

    return send_file(archivo, as_attachment=True)



@app.route('/premium')
def premium():
    return render_template('premium.html')

@app.route('/pro')
def pro():
    return render_template('pro.html')

@app.route('/pago-premium', methods=['GET', 'POST'])
def pago_premium():
    if "usuario_id" not in session:
        return redirect('/login')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        SELECT banco, titular, clabe, whatsapp, correo
        FROM configuracion_pago
        WHERE id=1
    """)
    datos = cursor.fetchone()

    if request.method == "POST":
        comprobante = request.files.get("comprobante")

        if comprobante and comprobante.filename != "":
            nombre_archivo = "premium_" + str(session["usuario_id"]) + "_" + comprobante.filename
            ruta = os.path.join("static", "comprobantes", nombre_archivo)
            comprobante.save(ruta)

            cursor.execute("""
                INSERT INTO pagos
                (usuario_id, nombre, correo, plan, monto, comprobante, estado, fecha)
                VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'))
            """, (
                session["usuario_id"],
                session["nombre"],
                session["correo"],
                "premium",
                "$99 MXN",
                nombre_archivo,
                "pendiente"
            ))

            conexion.commit()
            conexion.close()

            return render_template("pago_exitoso.html", plan="Premium")

    conexion.close()

    return render_template("pago_premium.html", datos=datos)

@app.route('/pago-pro', methods=['GET', 'POST'])
def pago_pro():
    if "usuario_id" not in session:
        return redirect('/login')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        SELECT banco, titular, clabe, whatsapp, correo
        FROM configuracion_pago
        WHERE id=1
    """)
    datos = cursor.fetchone()

    if request.method == "POST":
        comprobante = request.files.get("comprobante")

        if comprobante and comprobante.filename != "":
            nombre_archivo = "pro_" + str(session["usuario_id"]) + "_" + comprobante.filename
            ruta = os.path.join("static", "comprobantes", nombre_archivo)
            comprobante.save(ruta)

            cursor.execute("""
                INSERT INTO pagos
                (usuario_id, nombre, correo, plan, monto, comprobante, estado, fecha)
                VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'))
            """, (
                session["usuario_id"],
                session["nombre"],
                session["correo"],
                "pro",
                "$199 MXN",
                nombre_archivo,
                "pendiente"
            ))

            conexion.commit()
            conexion.close()

            return render_template("pago_exitoso.html", plan="Pro")

    conexion.close()

    return render_template("pago_pro.html", datos=datos)

@app.route('/registro', methods=['GET', 'POST'])
def registro():
    mensaje = ""

    if request.method == "POST":
        nombre = request.form.get("nombre")
        correo = request.form.get("correo")
        password = request.form.get("password")

        if not nombre or not correo or not password:
            mensaje = "Todos los campos son obligatorios."
            return render_template("registro.html", mensaje=mensaje)

        password_segura = generate_password_hash(password)

        try:
            conexion = sqlite3.connect("usuarios.db")
            cursor = conexion.cursor()

            cursor.execute("""
                INSERT INTO usuarios (nombre, correo, password, plan)
                VALUES (?, ?, ?, ?)
            """, (nombre, correo, password_segura, "gratis"))

            conexion.commit()
            conexion.close()

            return redirect('/login')

        except sqlite3.IntegrityError:
            mensaje = "Este correo ya está registrado."

        except Exception as e:
            mensaje = f"Error: {e}"

    return render_template("registro.html", mensaje=mensaje)
@app.route('/login', methods=['GET', 'POST'])
def login():
    mensaje = ""

    if request.method == "POST":
        correo = request.form.get("correo")
        password = request.form.get("password")

        conexion = sqlite3.connect("usuarios.db")
        cursor = conexion.cursor()

        cursor.execute("SELECT * FROM usuarios WHERE correo = ?", (correo,))
        usuario = cursor.fetchone()

        conexion.close()

        if usuario and check_password_hash(usuario[3], password):
            session["usuario_id"] = usuario[0]
            session["nombre"] = usuario[1]
            session["correo"] = usuario[2]
            session["plan"] = usuario[4]

            return redirect('/dashboard')
        else:
            mensaje = "Correo o contraseña incorrectos."

    return render_template("login.html", mensaje=mensaje)


@app.route('/dashboard')
def dashboard():
    if "usuario_id" not in session:
        return redirect('/login')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute(
        "SELECT nombre, correo, plan FROM usuarios WHERE id=?",
        (session["usuario_id"],)
    )

    usuario = cursor.fetchone()
    conexion.close()

    if usuario:
        session["nombre"] = usuario[0]
        session["correo"] = usuario[1]
        session["plan"] = usuario[2]

        return render_template(
            "dashboard.html",
            nombre=usuario[0],
            correo=usuario[1],
            plan=usuario[2]
        )

    return redirect('/login')

@app.route("/historial")
def historial():
    if "usuario_id" not in session:
        return redirect("/login")

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        SELECT nombre_archivo, resumen, total_palabras, fecha
        FROM documentos
        WHERE usuario_id = ?
        ORDER BY fecha DESC
    """, (session["usuario_id"],))

    documentos = cursor.fetchall()
    conexion.close()

    return render_template("historial.html", documentos=documentos)


@app.route('/logout')
def logout():
    session.clear()
    return redirect('/')

# ==========================
# PANEL ADMINISTRADOR PREMIUM
# ==========================

ADMIN_CLAVE = "admin123"


@app.route('/admin', methods=['GET', 'POST'])
def admin_login():
    mensaje = ""

    if request.method == "POST":
        clave = request.form.get("clave")

        if clave == ADMIN_CLAVE:
            session["admin"] = True
            return redirect('/admin/dashboard')
        else:
            mensaje = "Clave incorrecta."

    return render_template("admin_login.html", mensaje=mensaje)


@app.route('/admin/dashboard')
def admin_dashboard():
    if not session.get("admin"):
        return redirect('/admin')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("SELECT COUNT(*) FROM usuarios")
    total_usuarios = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM pagos")
    total_pagos = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM pagos WHERE estado='pendiente'")
    pagos_pendientes = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM usuarios WHERE plan='premium'")
    usuarios_premium = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM usuarios WHERE plan='pro'")
    usuarios_pro = cursor.fetchone()[0]

    conexion.close()

    return render_template(
        "admin_dashboard.html",
        total_usuarios=total_usuarios,
        total_pagos=total_pagos,
        pagos_pendientes=pagos_pendientes,
        usuarios_premium=usuarios_premium,
        usuarios_pro=usuarios_pro
    )


@app.route('/admin/usuarios')
def admin_usuarios():
    if not session.get("admin"):
        return redirect('/admin')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        SELECT id, nombre, correo, plan, fecha_registro
        FROM usuarios
        ORDER BY id DESC
    """)
    usuarios = cursor.fetchall()

    conexion.close()

    return render_template("admin_usuarios.html", usuarios=usuarios)


@app.route('/admin/usuarios/plan/<int:usuario_id>/<plan>')
def admin_cambiar_plan(usuario_id, plan):
    if not session.get("admin"):
        return redirect('/admin')

    if plan not in ["gratis", "premium", "pro"]:
        return redirect('/admin/usuarios')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute(
        "UPDATE usuarios SET plan=? WHERE id=?",
        (plan, usuario_id)
    )

    conexion.commit()
    conexion.close()

    return redirect('/admin/usuarios')


@app.route('/admin/pagos')
def admin_pagos():
    if not session.get("admin"):
        return redirect('/admin')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        SELECT id, nombre, correo, plan, monto, comprobante, estado, fecha, usuario_id
        FROM pagos
        ORDER BY id DESC
    """)
    pagos = cursor.fetchall()

    conexion.close()

    return render_template("admin_pagos.html", pagos=pagos)


@app.route('/admin/pagos/aprobar/<int:pago_id>/<int:usuario_id>/<plan>')
def admin_aprobar_pago(pago_id, usuario_id, plan):
    if not session.get("admin"):
        return redirect('/admin')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("UPDATE pagos SET estado='aprobado' WHERE id=?", (pago_id,))
    cursor.execute("UPDATE usuarios SET plan=? WHERE id=?", (plan, usuario_id))

    conexion.commit()
    conexion.close()

    return redirect('/admin/pagos')


@app.route('/admin/pagos/rechazar/<int:pago_id>')
def admin_rechazar_pago(pago_id):
    if not session.get("admin"):
        return redirect('/admin')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("UPDATE pagos SET estado='rechazado' WHERE id=?", (pago_id,))

    conexion.commit()
    conexion.close()

    return redirect('/admin/pagos')


@app.route('/admin/configuracion', methods=['GET', 'POST'])
def admin_configuracion():
    if not session.get("admin"):
        return redirect('/admin')

    conexion = sqlite3.connect("usuarios.db")
    cursor = conexion.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS configuracion_pago (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            banco TEXT,
            titular TEXT,
            clabe TEXT,
            whatsapp TEXT,
            correo TEXT
        )
    """)

    cursor.execute("SELECT COUNT(*) FROM configuracion_pago")
    existe = cursor.fetchone()[0]

    if existe == 0:
        cursor.execute("""
            INSERT INTO configuracion_pago
            (banco, titular, clabe, whatsapp, correo)
            VALUES (?, ?, ?, ?, ?)
        """, ("TU BANCO", "TU NOMBRE", "XXXX XXXX XXXX XXXX XX", "TU WHATSAPP", "contacto@studyai.com"))
        conexion.commit()

    if request.method == "POST":
        banco = request.form.get("banco")
        titular = request.form.get("titular")
        clabe = request.form.get("clabe")
        whatsapp = request.form.get("whatsapp")
        correo = request.form.get("correo")

        cursor.execute("""
            UPDATE configuracion_pago
            SET banco=?, titular=?, clabe=?, whatsapp=?, correo=?
            WHERE id=1
        """, (banco, titular, clabe, whatsapp, correo))

        conexion.commit()

    cursor.execute("""
        SELECT banco, titular, clabe, whatsapp, correo
        FROM configuracion_pago
        WHERE id=1
    """)

    datos = cursor.fetchone()
    conexion.close()

    return render_template("admin_configuracion.html", datos=datos)


@app.route('/admin/logout')
def admin_logout():
    session.pop("admin", None)
    return redirect('/admin')

@app.route('/recuperar', methods=['GET', 'POST'])
def recuperar():
    mensaje = ""

    if request.method == "POST":
        correo = request.form.get("correo")

        conexion = sqlite3.connect("usuarios.db")
        cursor = conexion.cursor()

        cursor.execute("SELECT id FROM usuarios WHERE correo=?", (correo,))
        usuario = cursor.fetchone()

        conexion.close()

        if usuario:
            session["correo_recuperar"] = correo
            return redirect('/restablecer')
        else:
            mensaje = "No existe una cuenta con ese correo."

    return render_template("recuperar.html", mensaje=mensaje)


@app.route('/restablecer', methods=['GET', 'POST'])
def restablecer():
    if "correo_recuperar" not in session:
        return redirect('/recuperar')

    mensaje = ""

    if request.method == "POST":
        nueva_password = request.form.get("password")
        confirmar = request.form.get("confirmar")

        if nueva_password != confirmar:
            mensaje = "Las contraseñas no coinciden."
        else:
            password_segura = generate_password_hash(nueva_password)

            conexion = sqlite3.connect("usuarios.db")
            cursor = conexion.cursor()

            cursor.execute(
                "UPDATE usuarios SET password=? WHERE correo=?",
                (password_segura, session["correo_recuperar"])
            )

            conexion.commit()
            conexion.close()

            session.pop("correo_recuperar", None)

            return redirect('/login')

    return render_template("restablecer.html", mensaje=mensaje)

@app.route("/politica")
def politica():

    return render_template("politica_privacidad.html")

@app.route("/terminos")
def terminos():

    return render_template("terminos.html")

@app.route("/cookies")
def cookies():

    return render_template("cookies.html")

if __name__ == "__main__":
    crear_base_datos()
    app.run(debug=True, port=5000)